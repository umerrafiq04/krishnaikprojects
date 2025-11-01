from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_rl_seminar_doc():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title with formatting
    title = doc.add_heading('Reinforcement Learning (RL)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('Seminar Report', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add decorative line
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_break()
    
    # Add summary section
    doc.add_heading('1. Brief Summary', level=1)
    summary_text = (
        "Reinforcement Learning (RL) is a machine learning paradigm where an agent learns to "
        "make optimal decisions by interacting with its environment. Through trial-and-error, "
        "the agent receives rewards or penalties and aims to maximize cumulative rewards over time. "
        "Key elements include:\n"
        "• Agent: The decision-maker\n"
        "• Environment: The world with which the agent interacts\n"
        "• State: Current situation\n"
        "• Action: Choices available\n"
        "• Reward: Feedback mechanism\n\n"
        "Popular algorithms: Q-Learning, Deep Q-Networks (DQN), Policy Gradients, PPO."
    )
    doc.add_paragraph(summary_text)
    
    # Add applications section with table
    doc.add_heading('2. Key Applications', level=1)
    applications = [
        ["Robotics", "Autonomous navigation, manipulation tasks, warehouse automation"],
        ["Gaming AI", "Mastering complex games (AlphaGo, StarCraft II)"],
        ["Autonomous Vehicles", "Real-time decision making in self-driving cars"],
        ["Recommendation Systems", "Personalized content delivery (Netflix, YouTube)"],
        ["Finance", "Algorithmic trading and portfolio optimization"],
        ["Healthcare", "Treatment planning and drug discovery"],
        ["Industrial Automation", "Energy optimization in smart grids"]
    ]
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'LightShading-Accent1'
    for app in applications:
        row = table.add_row().cells
        row[0].text = app[0]
        row[1].text = app[1]
    
    # Add scope section
    doc.add_heading('3. Scope', level=1)
    scope_text = (
        "RL is particularly effective for:\n"
        "✓ Sequential decision-making problems\n"
        "✓ Scenarios with delayed rewards\n"
        "✓ Complex environments without explicit supervision\n"
        "✓ Large state-space problems (enabled by deep RL)\n\n"
        "Current Limitations:\n"
        "• High computational requirements\n"
        "• Sample inefficiency (requires many trials)\n"
        "• Safety challenges during exploration\n"
        "• Difficulty in designing reward functions"
    )
    doc.add_paragraph(scope_text)
    
    # Add significance section
    doc.add_heading('4. Significance & Future Directions', level=1)
    doc.add_heading('Significance', level=2)
    sig_text = (
        "• Enables autonomy in unstructured environments\n"
        "• Solves problems intractable for traditional AI\n"
        "• Drives innovation across industries (healthcare, logistics)\n"
        "• Foundation for next-generation AI systems"
    )
    doc.add_paragraph(sig_text)
    
    doc.add_heading('Future Directions', level=2)
    future_text = (
        "• Improved Sample Efficiency: Meta-learning, transfer learning\n"
        "• Safe RL: Critical for medical and real-world applications\n"
        "• Multi-Agent Systems: Coordination in complex environments\n"
        "• Human-AI Collaboration: Learning from human feedback\n"
        "• Sim2Real Transfer: Bridging simulation-reality gap\n"
        "• Explainable RL: Interpretable decision-making"
    )
    doc.add_paragraph(future_text)
    
    # Add conclusion
    doc.add_heading('5. Conclusion', level=1)
    conclusion_text = (
        "Reinforcement Learning represents a transformative approach to artificial intelligence, "
        "enabling machines to learn complex behaviors through environmental interaction. "
        "Despite current challenges in efficiency and safety, ongoing advancements position RL "
        "as a cornerstone technology for future autonomous systems across diverse domains."
    )
    doc.add_paragraph(conclusion_text)
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "University Seminar | Reinforcement Learning Overview | Confidential"
    
    # Save document
    doc.save('Reinforcement_Learning_Seminar.docx')
    return "Word document created successfully!"

if __name__ == "_main_":
    create_rl_seminar_doc()
    print("Reinforcement_Learning_Seminar.docx created successfully!")